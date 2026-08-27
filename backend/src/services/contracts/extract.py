"""Trích text từ file hợp đồng và cắt thành từng điều khoản.

Cắt theo điều khoản (không theo số ký tự) để mỗi lần chấm rủi ro nhìn thấy trọn
một nghĩa vụ, và để trích dẫn ngược lại được đúng điều khoản trong hợp đồng.
"""
from __future__ import annotations

import logging
import re
import unicodedata
from dataclasses import dataclass

logger = logging.getLogger(__name__)

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TXT_MIME = "text/plain"

SUPPORTED_MIME_TYPES = {PDF_MIME, DOCX_MIME, TXT_MIME}

EXTENSION_TO_MIME = {".pdf": PDF_MIME, ".docx": DOCX_MIME, ".txt": TXT_MIME}

# Tiêu đề điều khoản hợp đồng Việt Nam thường gặp:
#   "Điều 5. Thanh toán", "ĐIỀU 5: THANH TOÁN", "Article 5.", "5. Thanh toán"
CLAUSE_HEADING_RE = re.compile(
    r"^\s*(?:(?:điều|article)\s+(\d+[a-zA-Z]?)|(\d{1,2}(?:\.\d{1,2})*)\.)\s*[\.\:\-–]?\s*(.{0,150})$",
    re.IGNORECASE,
)

# Điều khoản ngắn hơn ngưỡng này thường là tiêu đề mục hoặc dòng rác, ghép vào
# điều khoản trước thay vì tạo chunk riêng.
MIN_CLAUSE_CHARS = 80
# Cắt điều khoản quá dài để prompt không vượt context của model.
MAX_CLAUSE_CHARS = 6000


@dataclass(slots=True)
class Clause:
    """Một điều khoản hợp đồng đã tách."""

    position: int
    title: str | None
    text: str


class ExtractionError(RuntimeError):
    """File không đọc được hoặc không có text (ví dụ PDF scan ảnh)."""


def detect_mime(filename: str, declared: str | None) -> str:
    """Ưu tiên đuôi file vì browser hay khai sai ``application/octet-stream``."""

    suffix = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if suffix in EXTENSION_TO_MIME:
        return EXTENSION_TO_MIME[suffix]
    if declared in SUPPORTED_MIME_TYPES:
        return declared
    raise ExtractionError(
        f"Định dạng không hỗ trợ: {suffix or declared or 'không rõ'}. Chỉ nhận PDF, DOCX, TXT."
    )


def extract_text(payload: bytes, mime_type: str) -> str:
    """Trích toàn bộ text của file."""

    if mime_type == PDF_MIME:
        text = _extract_pdf(payload)
    elif mime_type == DOCX_MIME:
        text = _extract_docx(payload)
    elif mime_type == TXT_MIME:
        text = payload.decode("utf-8", errors="replace")
    else:
        raise ExtractionError(f"Định dạng không hỗ trợ: {mime_type}")

    text = _normalize(text)
    if len(text.strip()) < 50:
        raise ExtractionError(
            "Không trích được nội dung văn bản. Nếu đây là PDF scan từ ảnh, "
            "cần OCR trước khi tải lên."
        )
    return text


def _extract_pdf(payload: bytes) -> str:
    from io import BytesIO

    from pypdf import PdfReader

    try:
        reader = PdfReader(BytesIO(payload))
    except Exception as exc:
        raise ExtractionError(f"Không đọc được file PDF: {exc}") from exc
    if reader.is_encrypted:
        # Thử mở bằng mật khẩu rỗng: nhiều file chỉ đặt quyền chỉnh sửa.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ExtractionError("File PDF được đặt mật khẩu, không đọc được.") from exc
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(payload: bytes) -> str:
    from io import BytesIO

    import docx

    try:
        document = docx.Document(BytesIO(payload))
    except Exception as exc:
        raise ExtractionError(f"Không đọc được file DOCX: {exc}") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]
    # Điều khoản thanh toán/tiến độ hay nằm trong bảng, bỏ bảng là mất nội dung.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _normalize(text: str) -> str:
    """Chuẩn hóa Unicode và khoảng trắng, giữ lại cấu trúc dòng."""

    # NFC để "ệ" tổ hợp và "ệ" dựng sẵn so sánh được với nhau.
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    # Gộp nhiều dòng trống liên tiếp thành một.
    result: list[str] = []
    for line in lines:
        if not line and result and not result[-1]:
            continue
        result.append(line)
    return "\n".join(result).strip()


def split_clauses(text: str, *, max_clauses: int) -> list[Clause]:
    """Cắt hợp đồng thành danh sách điều khoản.

    Không tìm thấy tiêu đề điều khoản nào (hợp đồng viết liền một khối) thì
    quay về cắt theo đoạn văn để vẫn soát xét được.
    """

    blocks: list[tuple[str | None, list[str]]] = []
    for line in text.split("\n"):
        match = CLAUSE_HEADING_RE.match(line) if line.strip() else None
        if match and _looks_like_heading(line):
            number = match.group(1) or match.group(2)
            title = (match.group(3) or "").strip(" .:-–") or None
            heading = f"Điều {number}" + (f". {title}" if title else "")
            blocks.append((heading, [line]))
        elif blocks:
            blocks[-1][1].append(line)
        else:
            # Phần mở đầu (thông tin các bên) trước điều khoản đầu tiên.
            blocks.append((None, [line]))

    clauses = _finalize(blocks)
    if len(clauses) <= 1:
        clauses = _finalize(
            [(None, paragraph.split("\n")) for paragraph in text.split("\n\n") if paragraph.strip()]
        )

    if len(clauses) > max_clauses:
        logger.warning(
            "Hợp đồng có %d điều khoản, chỉ soát xét %d điều đầu", len(clauses), max_clauses
        )
        clauses = clauses[:max_clauses]
    return clauses


def _looks_like_heading(line: str) -> bool:
    """Loại bớt dương tính giả: dòng dài là câu văn, không phải tiêu đề."""

    stripped = line.strip()
    if len(stripped) > 160:
        return False
    # "5. Thanh toán" là tiêu đề, còn "5. Bên A phải thanh toán trong vòng..." thì
    # vẫn nhận vì dòng đủ ngắn; ngưỡng độ dài ở trên đã lọc phần lớn câu văn.
    return not stripped.endswith((",", ";"))


def _finalize(blocks: list[tuple[str | None, list[str]]]) -> list[Clause]:
    """Ghép các block quá ngắn rồi đánh số thứ tự."""

    merged: list[tuple[str | None, str]] = []
    for title, lines in blocks:
        body = "\n".join(lines).strip()
        if not body:
            continue
        if merged and len(body) < MIN_CLAUSE_CHARS:
            previous_title, previous_body = merged[-1]
            merged[-1] = (previous_title, f"{previous_body}\n{body}")
            continue
        merged.append((title, body))

    return [
        Clause(position=index, title=title, text=body[:MAX_CLAUSE_CHARS])
        for index, (title, body) in enumerate(merged, start=1)
    ]
